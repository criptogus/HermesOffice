#!/usr/bin/env python3
"""Generate templates/docs/invoice.docx — the HermesOffice invoice template.

House style (matches the other templates/docs/*.docx): English content,
[placeholder] tokens for the AI agents to replace, and a template note that
agents must delete before delivering. Regenerating this file is content-
deterministic (document parts are stable); the zip envelope bytes may vary
between runs, which is irrelevant to consumers.

Usage: env -u PYTHONPATH ~/.venvs/office/bin/python tools/gen-invoice-template.py
"""
from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
OUT = ROOT / "templates" / "docs" / "invoice.docx"

DARK = RGBColor(0x1F, 0x29, 0x37)  # near-black slate
ACCENT = RGBColor(0x6E, 0x4F, 0xF6)  # Hermes violet
GRAY = RGBColor(0x6B, 0x72, 0x80)
LIGHT_FILL = "F1F3F7"  # table header fill

PLACEHOLDER = "[...]"


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, inches: float) -> None:
    cell.width = Inches(inches)


def p(doc, text, *, size=10.5, bold=False, color=None, italic=False, space_after=4):
    par = doc.add_paragraph()
    par.paragraph_format.space_after = Pt(space_after)
    run = par.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color
    return par


def main() -> int:
    doc = Document()

    # base font
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)

    # margins
    for section in doc.sections:
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)

    # ── template note for AI agents (delete before delivering) ──
    note = doc.add_paragraph()
    note.paragraph_format.space_after = Pt(14)
    r = note.add_run(
        "Template note for AI agents: replace every [placeholder] with real "
        "content, keep the structure, recompute Subtotal/Discount/Tax/Total "
        "from the item rows, and delete this note before delivering."
    )
    r.font.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = GRAY

    # ── header ──
    p(doc, "[COMPANY NAME]", size=20, bold=True, color=DARK, space_after=0)
    p(doc, "[Street address, City, State ZIP] · [tax id if applicable]", size=9.5, color=GRAY, space_after=2)
    p(doc, "[billing@company.com] · [+1 555 000 0000]", size=9.5, color=GRAY, space_after=10)
    p(doc, "INVOICE", size=15, bold=True, color=ACCENT, space_after=8)

    # ── meta (label / value, borderless) ──
    meta = doc.add_table(rows=4, cols=2)
    meta.autofit = False
    meta_rows = [
        ("Invoice #", "[INV-YYYY-001]"),
        ("Date", "[YYYY-MM-DD]"),
        ("Due date", "[YYYY-MM-DD, e.g. net 30]"),
        ("Currency", "[USD | BRL | EUR]"),
    ]
    for i, (label, value) in enumerate(meta_rows):
        c0, c1 = meta.rows[i].cells
        set_cell_width(c0, 1.3)
        set_cell_width(c1, 3.5)
        c0.paragraphs[0].text = ""
        p0 = c0.paragraphs[0]
        p0.paragraph_format.space_after = Pt(3)
        rr = p0.add_run(label)
        rr.font.size = Pt(9.5)
        rr.font.bold = True
        rr.font.color.rgb = GRAY
        c1.paragraphs[0].text = ""
        p1 = c1.paragraphs[0]
        p1.paragraph_format.space_after = Pt(3)
        rr = p1.add_run(value)
        rr.font.size = Pt(10.5)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # ── bill to / from ──
    p(doc, "BILL TO", size=9.5, bold=True, color=GRAY, space_after=2)
    p(doc, "[Client company / full name]", size=11, bold=True, color=DARK, space_after=0)
    p(doc, "[Client street address, City, State ZIP]", size=10, color=DARK, space_after=0)
    p(doc, "[client@email.com]", size=10, color=DARK, space_after=10)

    # ── items table ──
    items = doc.add_table(rows=7, cols=4)
    items.style = doc.styles["Table Grid"]
    items.autofit = False
    headers = ["Description", "Qty", "Unit price", "Amount"]
    widths = [3.6, 0.7, 1.1, 1.2]
    for j, (h, w) in enumerate(zip(headers, widths)):
        cell = items.rows[0].cells[j]
        set_cell_width(cell, w)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.paragraphs[0].text = ""
        rr = cell.paragraphs[0].add_run(h)
        rr.font.size = Pt(9.5)
        rr.font.bold = True
        rr.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        if j >= 2:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        shade(cell, "1F2937")

    for row_idx in range(1, 6):  # 5 placeholder item rows
        cells = items.rows[row_idx].cells
        for j, w in enumerate(widths):
            set_cell_width(cells[j], w)
            cells[j].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cells[0].paragraphs[0].text = "[Item / service description]"
        cells[1].paragraphs[0].text = "[1]"
        cells[2].paragraphs[0].text = "[0.00]"
        cells[3].paragraphs[0].text = "[0.00]"
        for j in (1, 2, 3):
            cells[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for c in cells:
            for rr in c.paragraphs[0].runs:
                rr.font.size = Pt(9.5)

    # totals rows inside the items table (merged across the first 3 cols)
    totals = [
        ("Subtotal", "[0.00]"),
        ("Discount", "[0.00]"),
        ("Tax / VAT", "[0.00]"),
        ("TOTAL DUE", "[0.00]"),
    ]
    for idx, (label, value) in enumerate(totals):
        row = items.rows[5 + idx] if idx == 0 else items.add_row()
        a, b = row.cells[0], row.cells[3]
        merged = row.cells[0].merge(row.cells[1]).merge(row.cells[2])
        set_cell_width(merged, widths[0] + widths[1] + widths[2])
        set_cell_width(b, widths[3])
        merged.paragraphs[0].text = ""
        rr = merged.paragraphs[0].add_run(label)
        rr.font.size = Pt(10)
        rr.font.bold = label == "TOTAL DUE"
        b.paragraphs[0].text = ""
        rr = b.paragraphs[0].add_run(value)
        rr.font.size = Pt(10)
        rr.font.bold = label == "TOTAL DUE"
        b.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        if label == "TOTAL DUE":
            shade(merged, LIGHT_FILL)
            shade(b, LIGHT_FILL)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # ── payment terms + notes ──
    p(doc, "PAYMENT TERMS", size=9.5, bold=True, color=GRAY, space_after=2)
    p(doc, "[Method, bank details / IBAN / wire instructions]", size=10, space_after=6)
    p(doc, "NOTES", size=9.5, bold=True, color=GRAY, space_after=2)
    p(doc, "[Reference, PO number, or any additional context for the client]", size=10, space_after=6)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"written: {OUT}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
